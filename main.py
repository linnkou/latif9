from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    send_from_directory,
)
from docx import Document
from docx.oxml import parse_xml
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx.shared import Inches, RGBColor
import arabic_reshaper
from bidi.algorithm import get_display
import datetime
from openpyxl import load_workbook
import os


# دالة لتشكيل النصوص باللغة العربية
def reshape_text(text):
    return get_display(arabic_reshaper.reshape(text))


# دالة لتعديل اتجاه الكتابة من اليمين لليسار في مستند الـ Word
def set_right_to_left(doc):
    for paragraph in doc.paragraphs:
        paragraph.alignment = 3  # محاذاة النص إلى اليمين
        paragraph_xml = paragraph._element
        pPr = paragraph_xml.get_or_add_pPr()
        pPr.insert(
            0,
            parse_xml(
                r'<w:bidi xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            ),
        )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 3  # محاذاة النص في الخلايا إلى اليمين
                cell_xml = cell._element
                cell_pr = cell_xml.get_or_add_tcPr()
                cell_pr.insert(
                    0,
                    parse_xml(
                        r'<w:bidi xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    ),
                )


# دالة لإنشاء المخطط العمودي
def create_bar_chart(grades_count, class_name, output_path):
    if len(grades_count) == 0:
        return False

    categories = list(grades_count.keys())
    values = list(grades_count.values())
    reshaped_categories = [reshape_text(cat) for cat in categories]
    colors = sns.color_palette("muted", len(categories))

    fig, ax = plt.subplots(figsize=(12, 7))  # تكبير حجم الشكل
    bars = ax.bar(
        reshaped_categories, values, color=colors, edgecolor="black", linewidth=1.2
    )

    # تكبير حجم النصوص
    ax.set_xlabel(reshape_text("التقدير"), fontsize=16, labelpad=15)
    ax.set_ylabel(reshape_text("عدد التلاميذ"), fontsize=16, labelpad=15)
    ax.set_title(
        reshape_text(f"مخطط أعمدة يمثل عدد التلاميذ وتقديراتهم لقسم {class_name}"),
        fontsize=18,
        pad=20,
    )

    # إضافة القيم فوق الأعمدة بحجم نص أكبر
    for bar in bars:
        yval = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            yval + 0.5,
            f"{int(yval)}",
            ha="center",
            fontsize=14,
        )

    # تحسين شكل المحاور وتكبير حجم التسمية
    ax.tick_params(axis="x", labelsize=14)
    ax.tick_params(axis="y", labelsize=14)
    plt.xticks(rotation=45)  # تدوير تسميات المحور X لتجنب التداخل
    plt.tight_layout()
    plt.savefig(output_path, dpi=300)  # زيادة الدقة لتحسين وضوح النصوص
    plt.close()
    return True


# دالة لإنشاء المخطط الدائري
def create_pie_chart(grades_count, class_name, output_path):
    if len(grades_count) == 0:
        return False

    categories = list(grades_count.keys())
    values = list(grades_count.values())
    reshaped_categories = [reshape_text(cat) for cat in categories]
    colors = sns.color_palette("muted", len(categories))

    fig, ax = plt.subplots(figsize=(10, 10))  # تكبير حجم الشكل
    wedges, texts, autotexts = ax.pie(
        values,
        labels=reshaped_categories,
        autopct="%1.1f%%",
        startangle=90,
        colors=colors,
        textprops={"fontsize": 14},  # تكبير حجم النصوص
        pctdistance=0.85,
    )

    # تكبير حجم العنوان
    ax.set_title(
        reshape_text(f"مخطط دائري يمثل نسبة التلاميذ وتقديراتهم لقسم {class_name}"),
        fontsize=18,
        pad=20,
    )
    ax.set_aspect("equal")

    # تكبير حجم النصوص داخل المخطط
    for autotext in autotexts:
        autotext.set_fontsize(14)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300)  # زيادة الدقة لتحسين وضوح النصوص
    plt.close()
    return True


# دالة لتحليل البيانات وإنشاء التقرير
def process_file(file_path, num_classes, classes_data, grade_comments):
    try:
        xls = pd.ExcelFile(file_path)
        sheet_names = xls.sheet_names

        # فتح ملف Excel باستخدام openpyxl لتعديل الخلايا
        workbook = load_workbook(file_path)

        # إنشاء مستند Word
        doc = Document()
        doc.add_heading("تقرير نتائج التلاميذ", level=0).style.font.name = "Amiri"

        overall_data = []  # لتخزين البيانات العامة لكل قسم

        for i, sheet_name in enumerate(sheet_names):
            df = pd.read_excel(xls, sheet_name=sheet_name, skiprows=7)
            df.columns = (
                df.columns.astype(str).str.replace(" ", "").str.replace("/", "_")
            )
            required_columns = ["معدلتقويمالنشاطات_20", "الفرض_20", "الإختبار_20"]

            if df.empty or not all(col in df.columns for col in required_columns):
                continue

            df[required_columns] = df[required_columns].apply(
                pd.to_numeric, errors="coerce"
            )
            df.fillna(0, inplace=True)
            df["المعدل النهائي"] = (
                (df["معدلتقويمالنشاطات_20"] + df["الفرض_20"]) / 2
                + df["الإختبار_20"] * 2
            ) / 3

            def classify_grade(grade):  # تصنيف التقدير بناءً على الملاحظات المدخلة يدويًا
                if grade < 10:
                    return grade_comments[0]
                elif 10 <= grade < 12:
                    return grade_comments[1]
                elif 12 <= grade < 14:
                    return grade_comments[2]
                elif 14 <= grade < 16:
                    return grade_comments[3]
                elif 16 <= grade < 18:
                    return grade_comments[4]
                elif 18 <= grade <= 20:
                    return grade_comments[5]
                return "غير مصنف"

            df["التقدير"] = df["المعدل النهائي"].apply(classify_grade)
            class_name = classes_data[i]

            # إضافة الملاحظات في الخلايا H9 وما تحتها
            sheet = workbook[sheet_name]
            grades_count = df["التقدير"].value_counts().to_dict()

            for index, row in df.iterrows():
                cell_position = f"H{index + 9}"
                sheet[cell_position].value = row["التقدير"]

            doc.add_heading(f"نتائج قسم: {class_name}", level=1).style.font.name = (
                "Amiri"
            )
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = reshape_text("التقدير")
            hdr_cells[1].text = reshape_text("المعدل النهائي")
            hdr_cells[2].text = reshape_text("الاسم")
            hdr_cells[3].text = reshape_text("اللقب")
            hdr_cells[4].text = reshape_text("رقم التعريف")

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row["التقدير"]
                row_cells[1].text = f"{row['المعدل النهائي']:.2f}"
                row_cells[2].text = str(row.get("الاسم", "N/A"))
                row_cells[3].text = str(row.get("اللقب", "N/A"))
                row_cells[4].text = str(row.get("رقمالتعريف", "N/A"))

                if row["المعدل النهائي"] <= 10:
                    row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(
                        255, 0, 0
                    )

            # جمع بيانات القسم للإحصائيات النهائية
            highest_student = df.loc[df["المعدل النهائي"].idxmax()]
            lowest_student = df.loc[df["المعدل النهائي"].idxmin()]
            count_above_ten = (df["المعدل النهائي"] >= 10).sum()
            count_below_ten = (df["المعدل النهائي"] < 10).sum()

            overall_data.append(
                {
                    "class_name": class_name,
                    "highest_student": highest_student,
                    "lowest_student": lowest_student,
                    "count_above_ten": count_above_ten,
                    "count_below_ten": count_below_ten,
                    "total_students": len(df),
                }
            )

            bar_chart_path = f"static/bar_chart_{class_name}.png"
            pie_chart_path = f"static/pie_chart_{class_name}.png"

            os.makedirs("static", exist_ok=True)

            if create_bar_chart(grades_count, class_name, bar_chart_path):
                doc.add_paragraph(
                    f"المخطط العمودي للقسم: {class_name}"
                ).style.font.name = "Amiri"
                doc.add_picture(bar_chart_path, width=Inches(5))

            if create_pie_chart(grades_count, class_name, pie_chart_path):
                doc.add_paragraph(
                    f"المخطط الدائري للقسم: {class_name}"
                ).style.font.name = "Amiri"
                doc.add_picture(pie_chart_path, width=Inches(5))

            set_right_to_left(doc)

        # إضافة جدول الإحصائيات النهائية في نهاية التقرير
        doc.add_page_break()
        doc.add_heading("الإحصائيات النهائية", level=1).style.font.name = "Amiri"

        stats_table = doc.add_table(rows=1, cols=6)
        stats_table.style = "Table Grid"

        stats_hdr_cells = stats_table.rows[0].cells
        stats_hdr_cells[0].text = reshape_text("اسم القسم")
        stats_hdr_cells[1].text = reshape_text("صاحب أعلى معدل")
        stats_hdr_cells[2].text = reshape_text("صاحب أدنى معدل")
        stats_hdr_cells[3].text = reshape_text("عدد الطلاب (10 فما فوق)")
        stats_hdr_cells[4].text = reshape_text("عدد الطلاب (تحت 10)")
        stats_hdr_cells[5].text = reshape_text("النسبة المئوية (10 فما فوق)")

        for data in overall_data:
            row_cells = stats_table.add_row().cells

            row_cells[0].text = data["class_name"]

            highest_student_info = f"{data['highest_student']['الاسم']} {data['highest_student']['اللقب']} ({data['highest_student']['المعدل النهائي']:.2f})"
            lowest_student_info = f"{data['lowest_student']['الاسم']} {data['lowest_student']['اللقب']} ({data['lowest_student']['المعدل النهائي']:.2f})"

            row_cells[1].text = highest_student_info
            row_cells[2].text = lowest_student_info
            row_cells[3].text = str(data["count_above_ten"])
            row_cells[4].text = str(data["count_below_ten"])

            percentage_above_ten = (
                (data["count_above_ten"] / data["total_students"]) * 100
                if data["total_students"] > 0
                else 0
            )
            row_cells[5].text = f"{percentage_above_ten:.2f}%"

        set_right_to_left(doc)

        # حفظ التغييرات في ملف Excel
        workbook.save(file_path)

        report_name = f"static/تقرير_نتائج_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(report_name)

        return report_name

    except Exception as e:
        raise Exception(f"حدث خطأ: {str(e)}")


app = Flask(__name__, template_folder="templates", static_folder="static")
app.secret_key = "your_secret_key_here"  # Required for flash messages
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Disable caching for development

@app.route('/manifest.json')
def manifest():
    return send_from_directory('static', 'manifest.json')

@app.route('/sw.js')
def service_worker():
    return send_from_directory('static', 'sw.js')


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            # Handle form submission
            num_classes = request.form.get("num_classes", type=int)
            grade_comments = [request.form.get(f"comment_{i}") for i in range(6)]
            classes_data = [
                request.form.get(f"class_name_{i}") for i in range(num_classes)
            ]
            file = request.files.get("file")

            if file and num_classes > 0:
                os.makedirs("uploads", exist_ok=True)
                file_path = os.path.join("uploads", file.filename)
                file.save(file_path)

                report_path = process_file(
                    file_path, num_classes, classes_data, grade_comments
                )
                flash("تم إنشاء التقرير بنجاح", "success")
                return render_template(
                    "index.html",
                    excel_file=file_path,
                    word_file=report_path,
                    show_results=True
                )
            else:
                flash("يرجى إدخال عدد الأقسام وتحميل الملف", "error")
        except Exception as e:
            flash(str(e), "error")
        return redirect(url_for("index"))

    return render_template("index.html", show_results=False)


@app.route("/download/<file_type>/<path:filename>")
def download_file(file_type, filename):
    try:
        if file_type not in ["excel", "word"]:
            return "نوع ملف غير صالح", 400

        directory = "uploads" if file_type == "excel" else "static"
        return send_from_directory(
            directory=directory, path=filename, as_attachment=True
        )
    except Exception as e:
        flash(str(e), "error")
        return redirect(url_for("index"))


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
